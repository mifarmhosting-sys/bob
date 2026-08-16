from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = r"D:\gold-bot\Gold_Trading_Strategy_Video_Guide.docx"

BLUE = "2E74B5"; DARK = "1F4D78"; LIGHT = "E8EEF5"; PALE = "F4F6F9"; RED = "9B1C1C"; GOLD = "7A5A00"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    # Use transitional left/right names for compatibility with desktop Word.
    for m, v in [('top',top),('left',start),('bottom',bottom),('right',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(); fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar1,instr,fldChar2])

def add_bullet(doc, text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); p.add_run(text); return p

def add_num(doc, text):
    p=doc.add_paragraph(style='List Number'); p.add_run(text); return p

def callout(doc, title, text, color=PALE):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.5)
    c=t.cell(0,0); set_cell_shading(c,color); set_cell_margins(c,140,160,140,160)
    p=c.paragraphs[0]; r=p.add_run(title+'\n'); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK)
    p.add_run(text); doc.add_paragraph().paragraph_format.space_after=Pt(0)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor(0,0,0)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,DARK,10,5)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
for name in ['List Bullet','List Bullet 2','List Number']:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(11); s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.25
styles['List Bullet'].paragraph_format.left_indent=Inches(.375); styles['List Bullet'].paragraph_format.first_line_indent=Inches(-.188)
styles['List Number'].paragraph_format.left_indent=Inches(.375); styles['List Number'].paragraph_format.first_line_indent=Inches(-.188)

header=sec.header.paragraphs[0]; header.text='GOLD TRADING STRATEGY  |  VIDEO REFERENCE GUIDE'; header.style=styles['Normal']; header.runs[0].font.size=Pt(9); header.runs[0].font.color.rgb=RGBColor.from_string(DARK)
footer=sec.footer.paragraphs[0]; add_page_number(footer)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(60); p.paragraph_format.space_after=Pt(10)
r=p.add_run('Gold Trading Strategy'); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Chapter-by-Chapter Video Guide and Execution Process'); r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(12)
p.add_run('Based on “Gold Trading Strategy Live Explained by @HOLDwithPriyank”\nSatish K Finances • Video duration 55:30').italic=True
callout(doc,'Scope and accuracy note','This guide follows the video’s published chapter sequence and preserves the strategy rules, timings, and examples in faithful paraphrase. It is a structured study guide, not a verbatim transcript.',LIGHT)
callout(doc,'Risk warning','Educational material only—not financial or investment advice. Trading gold, futures, crypto-linked tokens, BTC, crude oil, or leveraged products can cause rapid and substantial losses. Verify instrument legality, regulation, fees, liquidity, and tax treatment for your jurisdiction; paper-trade before risking capital.', 'FCE8E6')
doc.add_page_break()

doc.add_heading('Quick Strategy Card', level=1)
tbl=doc.add_table(rows=1, cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.autofit=False
tbl.columns[0].width=Inches(1.875); tbl.columns[1].width=Inches(4.625)
hdr=tbl.rows[0].cells; hdr[0].text='Item'; hdr[1].text='Rule presented in the video'; set_repeat_table_header(tbl.rows[0])
for c in hdr: set_cell_shading(c,LIGHT); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
rows=[
('Chart preparation','Use the 15-minute chart to mark the reference session high/low; use the 5-minute chart for entry execution and finer confirmation.'),
('Asian reference range','Mark the Asian-session high and low at about 9:15–9:30 a.m. India time; do not trade inside that session range.'),
('London entry window','11:30 a.m.–2:30 p.m. India time.'),
('New York entry window','4:30 p.m.–7:30 p.m. India time, using the London-session range if it was sideways.'),
('Location filter','Trade only near the marked high or low; do not enter while price is in the middle of the range.'),
('Core event','Classify interaction with the level as a confirmed breakout or a fakeout; do not trust one candle alone.'),
('Direction','Confirmed upside breakout → long. Rejection/fakeout above the high → short. Mirror the logic at the low.'),
('Stop-loss','Beyond the invalidation swing/high/low, with a small buffer to reduce the chance of a news wick stopping the trade.'),
('Targets','Minimum 1:2 risk-to-reward. Main structural target is the opposite session boundary; breakout projection may use the range height.'),
('Management','At 2R/first target, book at least 60% and move the stop on the remainder to break-even; beginners may book 70–100%.'),
('Time exit','Hold intraday trades up to about 9:30 p.m.; if still awake, at most about 10:00–10:30 p.m., then exit before the next quiet Asian phase.'),
('Frequency','Aim for one carefully selected trade; avoid watching and trading the market 24 hours.'),
('Validation','Paper-trade 4–5 weeks and collect 20–25 trades before deciding whether the setup suits you.')]
for a,b in rows:
    cells=tbl.add_row().cells; cells[0].text=a; cells[1].text=b
    for c in cells: set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    cells[0].paragraphs[0].runs[0].bold=True

doc.add_heading('Execution Process: Exact Sequence', level=1)
steps=[
'Before the London window, open the 15-minute chart around 11:15 a.m. and mark the Asian-session high and low established by approximately 9:15–9:30 a.m.',
'Wait until 11:30 a.m. Do nothing while price remains in the middle of the marked range; attention begins only as price approaches or crosses a boundary.',
'Observe the level interaction. One candle is insufficient: use the following candle(s) to distinguish acceptance beyond the level from rejection back inside.',
'For a fakeout, switch to the 5-minute chart after confirmation and execute with the entry method you understand (price action, EMA, SMC/FVG/order block, VWAP, or candlestick logic).',
'Choose entry style: aggressive confirmation below/above the breakout candle’s extreme, or conservative confirmation through a nearby swing low/high (break of structure).',
'Place the stop beyond the invalidation swing, slightly buffered. Define risk before entering; do not widen the stop after entry.',
'Set the first target at 2R unless the opposite session boundary is closer; the opposite boundary is the main structural objective. For a breakout, the projected maximum target can equal the reference range height.',
'At the first target, book at least 60% and move the remaining stop to break-even. Let the remainder attempt the structural target.',
'If no Asian-range opportunity forms, mark the London high/low through 2:30 p.m. and repeat the same breakout/fakeout process in the 4:30–7:30 p.m. New York window.',
'Close remaining intraday exposure by about 9:30 p.m. (absolute late limit stated: roughly 10:00–10:30 p.m.). Record the trade and stop trading for the day.'
]
for s in steps: add_num(doc,s)

doc.add_heading('Decision Rules', level=1)
callout(doc,'Fakeout','Price pierces a marked boundary but closes back inside it. The next candle provides confirmation; the video repeatedly warns against acting on the first candle alone.',PALE)
callout(doc,'Confirmed breakout','A candle closes beyond the level; the next candle remains accepted beyond it and trades above/below the breakout candle’s extreme. The next one or two candles should not close back through the level.',PALE)
callout(doc,'No trade','Price is in the middle of the range, the entry window has passed, confirmation is unclear, the session has already made an unusually large move, or the trader cannot define a clean stop and 1:2 opportunity.','FFF4CE')

doc.add_page_break(); doc.add_heading('Chapter Guide', level=1)

chapters=[
('1. Teaser','00:00','The promise is a simple 1:2 framework, one trade a day, and chart examples intended to show the concept on real market days rather than only perfect historical selections.'),
('2. Introduction','00:21','The session is framed for beginners: how to plan gold entries, exits, and trade management.'),
('3. Can We Trade Gold in Today’s Time?','00:45','The speakers distinguish trading/buying gold from policy efforts to reduce physical imports. Instrument choice and regulatory context still matter.'),
('4. Government Bonds','01:15','Government gold-linked bonds are discussed as an alternative to physical imports, with fixed-return and appreciation concepts mentioned. Availability and terms must be independently verified.'),
('5. Gold ETFs','01:36','ETFs and a gold-linked token are discussed as access routes. The speaker warns that crypto tokens/exchanges may be legal to use yet unregulated, leaving counterparty and recovery risk.'),
('6. How Gold Became Volatile','02:52','Geopolitical uncertainty and wars are presented as drivers of higher volatility since roughly 2022–2024.'),
('7. Is Gold a Safe Asset Right Now?','04:04','Gold is described as a long-lived safe-haven asset, alongside land, while acknowledging that market conditions change.'),
('8. How to Start Trading Gold','04:30','The beginner setup is price-action based. It requires understanding a few candles and, most importantly, following fixed session timings rather than memorizing many patterns.'),
('9. The Problem of Overtrading','05:20','A 24-hour market tempts traders to stare at charts and trade repeatedly. Brokerage and poor decisions can turn a flat gross P&L into a net loss.'),
('10. Important Things to Remember','07:02','Work on India time, restrict participation to planned windows, and prioritize timing and discipline over constant activity.'),
('11. Understanding Different Trading Sessions','07:54','Asian range ends around 9:30 a.m.; London focus is 11:30 a.m.–2:30 p.m.; New York focus is 4:30–7:30 p.m. The video claims these six hours contain much of the day’s volume and movement.'),
('12. Breakout vs Fakeout','10:34','A breakout accepts price beyond a key level. A fakeout briefly breaches it and returns inside, trapping traders who entered immediately.'),
('13. Where to Draw Key Levels','16:18','Use objective session highs and lows instead of subjective support/resistance. These levels become the day’s decision points.'),
('14. When to Mark Highs and Lows?','19:01','On the 15-minute chart, mark the Asian high/low around the 9:15–9:30 a.m. cutoff, then wait until 11:30 a.m. before seeking a London-window entry.'),
('15. Which Strategy Should You Use?','21:03','The session framework is an overlay, not a compulsory indicator system. Entry confirmation may come from EMA, SMC/FVG/order blocks, VWAP, candlesticks, or another method the trader already understands.'),
('16. Live Trading Session','21:34','The example waits as price reaches the Asian high. A valid breakout can be traded long; a fakeout can be traded short. The opposite range edge is the main target, with at least 1:2 risk-to-reward.'),
('17. Different Entry Techniques','24:50','Two styles are separated: aggressive entry after a quicker candle confirmation and conservative entry after a swing break/break of structure.'),
('18. Aggressive Trade Entry','25:24','For a bearish fakeout above a level, mark the low of the breakout candle and enter only after another candle closes below it. Do not enter merely because price wicked through the level. Mirror for bullish setups.'),
('19. Conservative Trade Entry','26:25','Wait for a nearby swing low to break in a bearish setup (or swing high in a bullish setup). Enter on/after the confirming close; place the stop beyond the invalidation high/low with a small buffer.'),
('20. How to Manage a Trade?','28:29','At 2R/first target, book a minimum 60%. Beginners may book 70%, 80%, or all. Move the remaining stop to break-even and let the balance try for the opposite session boundary.'),
('21. Should You Close Trades After the London Session?','30:00','A London-window trade may be held while London/New York liquidity remains active. The stated intraday exit is about 9:30 p.m.; a late maximum of about 10:00–10:30 p.m. is mentioned.'),
('22. Liquidity Concept Explained','35:41','Liquidity is described as resting orders and stop-losses around obvious highs/lows. Price may sweep those pools before reversing; fakeouts exploit this behavior rather than treating every breach as continuation.'),
('23. What Do Different Trading Sessions Do?','38:30','Asian trade is presented as quieter and range-forming; London often expands or raids that range; New York may continue, reverse, or react to the London range. Session behavior provides context, not certainty.'),
('24. Which Trading Session Is Best?','44:53','The focus is London first, then New York if needed. Avoid weekend gold-token trading when volume is poor. Do not carry the previous day’s marked levels into a new daily process.'),
('25. Conclusion','51:03','Paper-trade for 4–5 weeks, gathering 20–25 observations. With fixed 1:2 risk/reward, the video illustrates that even a 50–60% win rate can remain positive before/after estimated fees; use your actual fees and slippage.')]
for title,ts,body in chapters:
    p=doc.add_heading(title,level=2); p.add_run('  ['+ts+']').italic=True
    doc.add_paragraph(body)

doc.add_page_break(); doc.add_heading('Entry and Management Playbook',level=1)
doc.add_heading('A. Bearish fakeout above a session high',level=2)
for x in ['Price trades above the marked high.','The candle closes back inside/below the range, and the following candle confirms rejection.','Aggressive: enter after a candle closes below the breakout candle’s low. Conservative: wait for a nearby swing low to break.','Stop above the rejection/swing high with a small buffer.','First target: 2R or the nearer structural objective. Main target: opposite session low.','At first target, book ≥60%; move the remainder to break-even.']: add_bullet(doc,x)
doc.add_heading('B. Bullish fakeout below a session low',level=2)
for x in ['Mirror the bearish rules: sweep below the low, close back inside, confirm, then enter after the breakout candle high or a swing high breaks.','Stop below the rejection/swing low with a buffer.','Target 2R and/or the opposite session high; manage partials identically.']: add_bullet(doc,x)
doc.add_heading('C. Confirmed breakout',level=2)
for x in ['First candle closes beyond the key level.','The next candle trades through the first candle’s extreme and remains accepted beyond the level; the next one or two candles should not close back through it.','Prefer a retest/known entry method rather than chasing.','Stop on the invalid side of the breakout structure.','Maximum projection discussed: one reference-range height from the breakout; still require acceptable risk-to-reward.']: add_bullet(doc,x)

doc.add_heading('Four-Week Validation Sheet',level=1)
tbl=doc.add_table(rows=1,cols=6); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.autofit=False
widths=[.8,.9,1.15,1.15,1.1,1.4]
for i,w in enumerate(widths): tbl.columns[i].width=Inches(w)
heads=['Date','Session','Setup','Entry style','Result (R)','Notes']
for i,h in enumerate(heads): tbl.cell(0,i).text=h; set_cell_shading(tbl.cell(0,i),LIGHT); set_cell_margins(tbl.cell(0,i)); tbl.cell(0,i).paragraphs[0].runs[0].bold=True
set_repeat_table_header(tbl.rows[0])
for _ in range(20):
    row=tbl.add_row();
    for c in row.cells: set_cell_margins(c,120,100,120,100)

doc.add_heading('Performance Math Used in the Video',level=1)
doc.add_paragraph('Illustration only: 20 trades, fixed ₹100 risk per trade, ₹200 target (1:2), and a rough fee estimate based on total turnover of wins and losses.')
for x in ['60% win rate: 12 wins = ₹2,400 gross profit; 8 losses = ₹800 gross loss; ₹1,600 before fees. The video subtracts an illustrative ₹320 fee estimate, leaving ₹1,280.','50% win rate: 10 wins = ₹2,000; 10 losses = ₹1,000; ₹1,000 before the same illustrative fee estimate.','Use actual brokerage, spread, slippage, funding, taxes, and instrument costs. A backtest or paper-trade result does not guarantee live performance.']: add_bullet(doc,x)

doc.add_heading('Final Pre-Trade Checklist',level=1)
checks=['Today is a liquid weekday; no major scheduled news is being ignored.','Asian or London high/low is marked from the correct session and day.','Current time is inside 11:30–2:30 or 4:30–7:30 India time.','Price is at a boundary—not in the middle of the range.','Breakout/fakeout confirmation uses more than one candle.','Entry method and invalidation point are clear.','Stop and position size keep risk within the personal limit.','At least 1:2 is available, or the structural target is closer and the trade is skipped/adjusted.','Partial-profit and break-even rules are entered before the trade.','No second impulsive trade after the planned opportunity.']
for x in checks: add_bullet(doc,'☐ '+x)

doc.add_heading('Source',level=1)
p=doc.add_paragraph(); p.add_run('Video: ').bold=True; p.add_run('https://www.youtube.com/watch?v=qB8uKV5qkX4&t=1494s')
doc.add_paragraph('Chapter timestamps are taken from the video description/transcript as displayed on YouTube. Strategy descriptions are paraphrased for study and operational clarity.')

doc.core_properties.title='Gold Trading Strategy — Video Guide'
doc.core_properties.subject='Chaptered study guide based on the referenced YouTube video'
doc.core_properties.author='Prepared for the user'
doc.save(OUT)
print(OUT)
